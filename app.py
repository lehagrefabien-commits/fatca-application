import os
import uuid
from datetime import datetime

from flask import Flask, render_template, request, send_file, abort, url_for
from docx import Document

from pathlib import Path
from threading import Lock

# Dossier racine du projet (là où se trouve app.py)
APP_DIR = os.path.dirname(os.path.abspath(__file__))

# =========================
# Compteur de formulaires générés
# =========================
COUNTER_FILE = Path(APP_DIR) / "counter.txt"
COUNTER_LOCK = Lock()

def get_counter() -> int:
    if not COUNTER_FILE.exists():
        COUNTER_FILE.write_text("0", encoding="utf-8")
        return 0
    try:
        return int(COUNTER_FILE.read_text(encoding="utf-8").strip())
    except ValueError:
        return 0

def increment_counter() -> int:
    with COUNTER_LOCK:
        value = get_counter() + 1
        COUNTER_FILE.write_text(str(value), encoding="utf-8")
        return value

# Dossier de sortie pour les fichiers générés
OUTPUT_DIR = os.path.join(APP_DIR, "output")
os.makedirs(OUTPUT_DIR, exist_ok=True)

# Application Flask (c'est CE "app" que Gunicorn doit importer)
app = Flask(__name__)

# =========================
# 1) Contexte (page d'accueil)
# =========================
CONTEXT = {
    "fr": (
        "<strong>Accord FATCA – Contexte juridique</strong><br><br>"
        "Par un arrêt interlocutoire rendu fin 2025, la Cour des marchés de Belgique a décidé de saisir "
        "la Cour de justice de l’Union européenne (CJUE) de treize questions préjudicielles "
        "relatives à la conformité de l’accord intergouvernemental FATCA avec le droit de "
        "l’Union européenne, et en particulier avec les exigences du règlement général sur "
        "la protection des données (RGPD).<br><br>"
        "Cette saisine intervient à la suite de la décision n°79/2025 du 24 avril 2025 de "
        "l’Autorité de protection des données (APD), par laquelle plusieurs violations du "
        "RGPD ont été constatées. Cette décision impose au SPF Finances, en tant que "
        "responsable de traitement, de mettre en conformité les transferts de données "
        "opérés dans le cadre de l’accord FATCA afin de les rendre compatibles avec le RGPD.<br><br>"
        "Dans ce contexte juridique désormais porté au niveau européen, les personnes "
        "concernées par l’accord FATCA disposent du droit de demander l’effacement de leurs "
        "données à caractère personnel lorsqu’elles estiment que leur traitement ne respecte "
        "pas le RGPD.<br><br>"
        "La présente démarche a pour objet de vous permettre d’exercer concrètement ce "
        "droit à l’effacement auprès du SPF Finances (l’administration fiscale belge), "
        "en générant un courrier prêt à être imprimé, signé et adressé."
    ),
    "nl": (
        "<strong>Context – FATCA-akkoord</strong><br><br>"
        "Bij een interlocutoir arrest dat eind 2025 werd gewezen, heeft het Marktenhof van België "
        "beslist om dertien prejudiciële vragen voor te leggen aan het Hof van Justitie van "
        "de Europese Unie (HvJ-EU) met betrekking tot de verenigbaarheid van het "
        "intergouvernementele FATCA-akkoord met het recht van de Europese Unie, en in het "
        "bijzonder met de vereisten van de Algemene Verordening Gegevensbescherming (AVG).<br><br>"
        "Deze verwijzing volgt op beslissing nr. 79/2025 van 24 april 2025 van de "
        "Gegevensbeschermingsautoriteit (GBA), waarin meerdere schendingen van de AVG werden "
        "vastgesteld. In deze beslissing wordt de FOD Financiën, in zijn hoedanigheid van "
        "verwerkingsverantwoordelijke, verplicht om de gegevensoverdrachten die plaatsvinden "
        "in het kader van het FATCA-akkoord in overeenstemming te brengen met de AVG.<br><br>"
        "In deze juridische context, die inmiddels op Europees niveau wordt behandeld, "
        "beschikken personen die door het FATCA-akkoord worden getroffen over het recht om "
        "de wissing van hun persoonsgegevens te verzoeken wanneer zij van oordeel zijn dat "
        "de verwerking ervan niet in overeenstemming is met de AVG.<br><br>"
        "Het doel van deze toepassing is u in staat te stellen dit recht op gegevenswissing "
        "concreet uit te oefenen ten aanzien van de FOD Financiën (de Belgische "
        "belastingadministratie), door het genereren van een brief die kan worden afgedrukt, "
        "ondertekend en verzonden."
    ),
}

# =========================
# 2) Textes UI (formulaire & page merci)
# =========================
TEXTS = {
    "fr": {
        "title": "Générer une demande d’effacement FATCA",
        "subtitle": "Renseignez vos coordonnées pour générer un document Word prêt à imprimer et signer.",
        "civilite": "Civilité",
        "monsieur": "Monsieur",
        "madame": "Madame",
        "prenom": "Prénom",
        "nom": "Nom",
        "adresse": "Adresse (rue, numéro)",
        "code_postal": "Code postal",
        "ville": "Ville",
        "pays": "Pays",
        "date_naissance": "Date de naissance",
        "date_naissance_ph": "JJ/MM/AAAA",
        "ville_naissance": "Ville de naissance",
        "pays_naissance": "Pays de naissance",
        "generate": "Générer le document",
        "privacy": "Aucune donnée n’est conservée pour la génération du document : les informations saisies sur ce formulaire servent uniquement à produire le courrier.",
        "ready_title": "Votre document est prêt",
        "ready_sub": "Téléchargez, imprimez et signez votre courrier avant de l’envoyer au SPF Finances.",
        "download": "Télécharger le document",
        "support_title": "Soutenir l’Association des Américains Accidentels (AAA)",
        "support_sub": "Si ce générateur vous a été utile, vous pouvez soutenir l’action de l’AAA via le formulaire ci-dessous.",
    },
    "nl": {
        "title": "FATCA-verwijderingsverzoek genereren",
        "subtitle": "Vul uw gegevens in om een Word-document te genereren dat u kunt afdrukken en ondertekenen.",
        "civilite": "Aanspreking",
        "monsieur": "Heer",
        "madame": "Mevrouw",
        "prenom": "Voornaam",
        "nom": "Familienaam",
        "adresse": "Adres (straat, nummer)",
        "code_postal": "Postcode",
        "ville": "Gemeente",
        "pays": "Land",
        "date_naissance": "Geboortedatum",
        "date_naissance_ph": "DD/MM/JJJJ",
        "ville_naissance": "Geboorteplaats",
        "pays_naissance": "Geboorteland",
        "generate": "Document genereren",
        "privacy": "Er worden geen persoonsgegevens bewaard voor de documentgeneratie: de ingevoerde gegevens worden enkel gebruikt om de brief te genereren.",
        "ready_title": "Uw document is klaar",
        "ready_sub": "Download, druk af en onderteken uw brief voordat u die naar de FOD Financiën verstuurt.",
        "download": "Document downloaden",
        "support_title": "Steun de Association des Américains Accidentels (AAA)",
        "support_sub": "Als deze tool nuttig was, kunt u de werking van AAA steunen via het formulier hieronder.",
    },
}

# =========================
# 2bis) Textes UI (page Contexte)
# =========================
CONTEXT_UI = {
    "fr": {
        "page_title": "Accord FATCA – Contexte juridique",
        "kpi_label": "courriers déjà générés",
        "kpi_note": "Indicateur mis à jour automatiquement.",
        "kpi_desc": (
            "Ce chiffre correspond au nombre de courriers générés via cette application afin d’exercer "
            "le droit d’accès et le droit à l’effacement des données personnelles, tels que garantis par le RGPD, "
            "dans le contexte de l’accord FATCA."
        ),
        "cta_primary": "Continuer →",
        "cta_secondary": "← Retour",
    },
    "nl": {
        "page_title": "FATCA-akkoord – Juridische context",
        "kpi_label": "brieven reeds gegenereerd",
        "kpi_note": "Indicator wordt automatisch bijgewerkt.",
        "kpi_desc": (
            "Dit cijfer geeft het aantal brieven weer dat via deze toepassing werd gegenereerd om het recht op inzage "
            "en het recht op gegevenswissing uit te oefenen, zoals gewaarborgd door de AVG, in het kader van het FATCA-akkoord."
        ),
        "cta_primary": "Verder →",
        "cta_secondary": "← Terug",
    },
}

# =========================
# 3) Remplacement dans Word
# =========================
def _replace_in_paragraph(paragraph, mapping: dict):
    full_text = "".join(run.text for run in paragraph.runs)
    changed = False
    for k, v in mapping.items():
        if k in full_text:
            full_text = full_text.replace(k, v)
            changed = True

    if changed:
        if paragraph.runs:
            for run in paragraph.runs:
                run.text = ""
            paragraph.runs[0].text = full_text
        else:
            paragraph.add_run(full_text)

def _replace_in_doc(doc: Document, mapping: dict):
    for p in doc.paragraphs:
        _replace_in_paragraph(p, mapping)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    _replace_in_paragraph(p, mapping)

# =========================
# 4) Routes
# =========================
@app.route("/", methods=["GET"])
def index():
    return render_template("lang_select.html")

@app.route("/context/<lang>", methods=["GET"])
def context(lang):
    if lang not in ("fr", "nl"):
        abort(404)

    count = get_counter()

    return render_template(
        "langue.html",
        context_text=CONTEXT[lang],     # ton HTML FR/NL existant
        lang=lang,                     # la langue de la page
        t=CONTEXT_UI[lang],            # textes UI FR/NL pour l'encart compteur + titres/boutons
        generated_count=count,         # compteur affiché sur la page Contexte uniquement
    )

@app.route("/fr", methods=["GET"])
def form_fr():
    return _render_form("fr")

@app.route("/nl", methods=["GET"])
def form_nl():
    return _render_form("nl")

def _render_form(lang: str):
    if lang not in ("fr", "nl"):
        abort(404)

    return render_template(
        "form.html",
        lang=lang,
        t=TEXTS[lang],
    )

@app.route("/generate", methods=["POST"])
def generate():
    increment_counter()
    lang = request.form.get("lang", "fr").strip()
    if lang not in ("fr", "nl"):
        lang = "fr"

    civilite = request.form.get("civilite", "H").strip()  # H / F
    prenom = request.form.get("prenom", "").strip()
    nom = request.form.get("nom", "").strip()
    adresse = request.form.get("adresse", "").strip()
    code_postal = request.form.get("code_postal", "").strip()
    ville = request.form.get("ville", "").strip()
    pays = request.form.get("pays", "").strip()
    date_naissance = request.form.get("date_naissance", "").strip()
    ville_naissance = request.form.get("ville_naissance", "").strip()
    pays_naissance = request.form.get("pays_naissance", "").strip()

    if not all([prenom, nom, adresse, code_postal, ville, pays, date_naissance, ville_naissance, pays_naissance]):
        abort(400, "Champs manquants. Merci de compléter tous les champs.")

    cp_ville = f"{code_postal} {ville}"

    template_filename = "template_fr.docx" if lang == "fr" else "template_nl.docx"
    template_path = os.path.join(APP_DIR, "templates", "word_templates", template_filename)
    if not os.path.exists(template_path):
        abort(500, f"Template Word introuvable : templates/word_templates/{template_filename}")

    date_courrier = datetime.now().strftime("%d/%m/%Y")

    # --- Variables attendues par les nouveaux templates Word ---
    mapping = {
        "{{PRENOM}}": prenom,
        "{{NOM}}": nom,
        "{{ADRESSE}}": adresse,
        "{{CP_VILLE}}": cp_ville,
        "{{PAYS}}": pays,
        "{{VILLE}}": ville,
        "{{DATE}}": date_courrier,
        "{{DATE_NAISSANCE}}": date_naissance,
        "{{VILLE_NAISSANCE}}": ville_naissance,
        "{{PAYS_NAISSANCE}}": pays_naissance,
    }

    # --- Gestion du "né / née" en FR + valeur NL ---
    if lang == "fr":
        mapping["{{NE}}"] = "née" if civilite == "F" else "né"
    else:
        # NL: forme neutre standard
        mapping["{{NE}}"] = "geboren"

    doc = Document(template_path)
    _replace_in_doc(doc, mapping)

    token = str(uuid.uuid4())
    output_path = os.path.join(OUTPUT_DIR, f"{token}.docx")
    doc.save(output_path)

    download_url = url_for("download", token=token)
    return render_template(
    "merci.html",
    download_url=download_url,
    t=TEXTS[lang],
    lang=lang
)


@app.route("/download/<token>", methods=["GET"])
def download(token: str):
    path = os.path.join(OUTPUT_DIR, f"{token}.docx")
    if not os.path.exists(path):
        abort(404)

    return send_file(
        path,
        as_attachment=True,
        download_name="Demande_effacement_FATCA.docx",
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )
